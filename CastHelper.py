from PIL import ImageTk, Image
import tkinter as tk
import helper as helper

from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION

"""
TODO:
1.	Place necessaire 
2.	Mesurer diameter MD des 12 dents
3.	Centre molaire et PM
4.	Plan sagitale median

"""


'''
Points:

[L0_sella,
L1_nasion,
L2_porion,
L3_orbitale,
L4_ENP,
L5_ENA,
L6_point_A,
L7_point_B,
L8_pogonion,
L9_menton,
L10_gnathion,
L11_symphyse
L12_gonion,
L13_articulate,
L14_I1,
L15_I2,
L16_i1,
L17_i2,
L18_M1,
L19_M2,
L20_m1,
L21_m2,
L22_POA,
L23_POP]
'''

class Cephalo(tk.Tk):
    def __init__(self):
        tk.Tk.__init__(self)
        self.name = "Boursas Yahia"
        self.x = self.y = 0
        self.pts = []
        self.pts_text = ["ncs1", "ncs2", "ncs3", "ncs4", "ncs5", "d_32", "m32","d_31", "m31",
                         "m_41", "d41","m_42", "d42", "d36", "m36", "clb1", "clb2"]

        self.angles_text = ["Espace necessaire"]

        self.canvas = tk.Canvas(self, width=1000, height=750, cursor="cross")
        self.canvas.pack(side="top", fill="both", expand=True)
        # self.img = ImageTk.PhotoImage(Image.open(self.name+".jpg"))
        self.img = ImageTk.PhotoImage(Image.open("cast.jpg"))
        self.canvas.create_image(0, 0, image=self.img, anchor='nw')
        self.canvas.bind("<ButtonPress-1>", self.on_button_press)
        self.canvas.bind("<ButtonPress-3>", self.on_back)

    def on_back(self, event):
        print(self.pts_text[len(self.pts) - 1])
        self.canvas.delete(self.pts_text[len(self.pts) - 1])
        self.canvas.delete(str(self.pts_text[len(self.pts) - 1]) + "_pt")
        self.pts.pop(-1)

#     def calculate(self):
#
#         self.angles = [SNA, SNB, ANB, AOBO, SND, AC, AF, FMA, AG,
#                        AXE_Y, E_SUP, E_INF, I_F, I_M, I_I, ALPHA, BETA]
#
#         self.description = [SNA_dsc, SNB_dsc, ANB_dsc, AOBO_dsc, SND_dsc, AC_dsc,
#                             AF_dsc, FMA_dsc, AG_dsc, AXE_Y_dsc, E_SUP_dsc, E_INF_dsc,
#                             I_F_dsc, I_M_dsc, I_I_dsc, ALPHA_dsc, BETA_dsc]
#
# # //////////////////////////////////
#         document = Document()
#         # section = document.sections[-1]
#         # new_width, new_height = section.page_height, section.page_width
#         # new_section = document.add_section(WD_SECTION.NEW_PAGE)
#         # new_section.orientation = WD_ORIENT.LANDSCAPE
#         # new_section.page_width = new_width
#         # new_section.page_height = new_height
#
#         document.add_heading(self.name, 0)
# # ////////////////////////////////
#
#         for i, angle in enumerate(self.angles):
#             self.canvas.create_text(120, 20 * i + 500, fill="black", font="Times 12 bold",
#                                     text=self.angles_text[i] + ": " + str(angle))
#             p = self.angles_text[i] + ": " + str(angle) + ", " + self.description[i]
#             document.add_paragraph(p, style='List Bullet')
#         document.add_picture(self.name + '.jpg', width=Inches(4))
#         document.save(self.name + ".docx")

    def calculate(self):

        disponible = helper.distance(self.pts[0], self.pts[1], self.pts[2], self.pts[3],
                                     self.pts[4], self.pts[-1], self.pts[-2])

        la_32 = helper.distance2(self.pts[5], self.pts[6], self.pts[-1], self.pts[-2])
        la_31 = helper.distance2(self.pts[7], self.pts[8], self.pts[-1], self.pts[-2])
        la_41 = helper.distance2(self.pts[9], self.pts[10], self.pts[-1], self.pts[-2])
        la_42 = helper.distance2(self.pts[11], self.pts[12], self.pts[-1], self.pts[-2])
        la_36 = helper.distance2(self.pts[13], self.pts[14], self.pts[-1], self.pts[-2])

        print("la32  la31  la41  la42  la36  disponible")
        print(la_32, la_31, la_41, la_42, disponible)

        p10 = (la_31+la_36)*3.85
        arc = (la_42+ la_41+ la_31+ la_32)
        necessaire = 2*arc + 21
        ddm = (disponible-3.4) - necessaire

        print(p10, arc, necessaire, disponible, ddm)
        # self.canvas.create_text(120, 20*1+500, fill="black", font="Times 12 bold",
        #                         text=str(necessaire))


    def on_button_press(self, event):
        # 26
        if len(self.pts) == 17:
            self.calculate()
        else:
            self.x = event.x
            self.y = event.y
            x0, y0 = (self.x - 2, self.y - 2)
            x1, y1 = (event.x + 2, event.y + 2)
            self.canvas.create_oval(x0, y0, x1, y1, fill="red", tag=str(self.pts_text[len(self.pts)]) + "_pt")
            self.canvas.create_text(self.x, self.y - 12, fill="red", font="Times 12 bold",
                                    text=self.pts_text[len(self.pts)], tag=self.pts_text[len(self.pts)])
            self.pts.append([float(self.x), float(self.y)])

            # if len(self.pts) == 3:
            #     self.angle = helper.get_angle(self.pts[0], self.pts[1], self.pts[2])
            #     self.canvas.create_text(self.pts[1][0], self.pts[1][1]-12, fill="red", font="Times 12",
            #                              text=self.angle)


if __name__ == "__main__":
    app = Cephalo()
    app.mainloop()






